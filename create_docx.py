import os
import sys
import re
import html

try:
    import docx
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def add_hyperlink(paragraph, url, text, color="0563C1", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = parse_xml(
        f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:id="{r_id}"/>'
    )
    new_run = parse_xml('<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    rPr = parse_xml('<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    
    rPr.append(parse_xml('<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:ascii="Calibri" w:hAnsi="Calibri"/>'))
    rPr.append(parse_xml('<w:sz xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="21"/>'))

    if color:
        c = parse_xml(f'<w:color xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="{color}"/>')
        rPr.append(c)

    if underline:
        u = parse_xml('<w:u xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="single"/>')
        rPr.append(u)

    new_run.append(rPr)
    new_run_text = parse_xml(f'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">{html.escape(text)}</w:t>')
    new_run.append(new_run_text)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def add_formatted_runs(paragraph, text, base_font_size=10.5, font_name="Calibri"):
    # Matches: links [text](url), bold **text**, code `text`, italic *text*, raw URLs http(s)://...
    pattern = re.compile(
        r'(\[([^\]]+)\]\(([^)]+)\))|(\*\*([^*]+)\*\*)|(`([^`]+)`)|(\*([^*]+)\*)|(https?://[^\s]+)'
    )
    pos = 0
    for m in pattern.finditer(text):
        start, end = m.span()
        if start > pos:
            run = paragraph.add_run(text[pos:start])
            run.font.name = font_name
            run.font.size = Pt(base_font_size)
            run.font.color.rgb = RGBColor(38, 38, 38)
        
        if m.group(1):  # Link [text](url)
            link_text = m.group(2).strip('`')
            url = m.group(3)
            add_hyperlink(paragraph, url, link_text)
        elif m.group(4):  # Bold **text**
            run = paragraph.add_run(m.group(5))
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(base_font_size)
            run.font.color.rgb = RGBColor(38, 38, 38)
        elif m.group(6):  # Code `text`
            run = paragraph.add_run(m.group(7))
            run.font.name = "Consolas"
            run.font.size = Pt(base_font_size - 1.0)
            run.font.color.rgb = RGBColor(160, 30, 30)
        elif m.group(8):  # Italic *text*
            run = paragraph.add_run(m.group(9))
            run.italic = True
            run.font.name = font_name
            run.font.size = Pt(base_font_size)
            run.font.color.rgb = RGBColor(38, 38, 38)
        elif m.group(10): # Raw URL
            raw_url = m.group(10)
            add_hyperlink(paragraph, raw_url, raw_url)

        pos = end

    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = font_name
        run.font.size = Pt(base_font_size)
        run.font.color.rgb = RGBColor(38, 38, 38)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="none"/><w:left w:val="none"/><w:right w:val="none"/></w:tblBorders>'
        )
        tblPr[0].append(borders)

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pBdr = parse_xml('<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
    p._p.get_or_add_pPr().append(pBdr)

def parse_markdown_to_docx(md_filepath, docx_filepath):
    with open(md_filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Set margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    base_dir = os.path.dirname(os.path.abspath(md_filepath))

    i = 0
    n = len(lines)

    while i < n:
        line = lines[i].rstrip('\r\n')
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Horizontal Rule
        if line.strip() == '---':
            add_divider(doc)
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            title_text = line[2:].strip()
            run = p.add_run(title_text)
            run.font.name = "Calibri"
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(27, 54, 93) # Deep Navy
            i += 1
            continue

        if line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            heading_text = line[3:].strip()
            add_formatted_runs(p, heading_text, base_font_size=14, font_name="Calibri")
            # Override heading color to Deep Navy
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(27, 54, 93)
            i += 1
            continue

        if line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            heading_text = line[4:].strip()
            add_formatted_runs(p, heading_text, base_font_size=12, font_name="Calibri")
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(46, 117, 182) # Steel Blue
            i += 1
            continue

        # Image block: ![alt](./path)
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', line.strip())
        if img_match:
            alt_text = img_match.group(1)
            img_rel_path = img_match.group(2)
            img_full_path = os.path.normpath(os.path.join(base_dir, img_rel_path))
            
            if os.path.exists(img_full_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run()
                run.add_picture(img_full_path, width=Inches(6.2))
            i += 1
            continue

        # Table block
        if line.strip().startswith('|'):
            table_lines = []
            while i < n and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Filter out separator row (|---|---|)
            rows_data = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                if all(re.match(r'^:?-+:?$', c) for c in cells):
                    continue
                rows_data.append(cells)
            
            if rows_data:
                num_cols = max(len(r) for r in rows_data)
                num_rows = len(rows_data)

                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(table)

                col_widths = [Inches(1.5), Inches(2.35), Inches(2.35)] if num_cols == 3 else [Inches(6.2 / num_cols)] * num_cols

                for r_idx, row_cells_data in enumerate(rows_data):
                    row = table.rows[r_idx]
                    is_header = (r_idx == 0)

                    trPr = row._element.get_or_add_trPr()
                    trPr.append(parse_xml('<w:cantSplit xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

                    if is_header:
                        trPr.append(parse_xml('<w:tblHeader xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

                    for c_idx in range(num_cols):
                        cell = row.cells[c_idx]
                        cell.width = col_widths[c_idx]
                        set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
                        
                        if is_header:
                            set_cell_background(cell, "1B365D") # Deep navy header
                        else:
                            fill_color = "F7F9FA" if r_idx % 2 == 1 else "FFFFFF" # Alternating row colors
                            set_cell_background(cell, fill_color)
                        
                        cell_text = row_cells_data[c_idx] if c_idx < len(row_cells_data) else ""
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.line_spacing = 1.15
                        
                        add_formatted_runs(p, cell_text, base_font_size=9.5, font_name="Calibri")
                        
                        if is_header:
                            for r in p.runs:
                                r.bold = True
                                r.font.color.rgb = RGBColor(255, 255, 255)

                p_space = doc.add_paragraph()
                p_space.paragraph_format.space_before = Pt(0)
                p_space.paragraph_format.space_after = Pt(6)

            continue

        # Bullet / Numbered list items
        bullet_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)$', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            marker = bullet_match.group(2)
            content = bullet_match.group(3)

            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)

            level = indent // 2
            if level == 0:
                p.paragraph_format.left_indent = Inches(0.25)
            elif level == 1:
                p.paragraph_format.left_indent = Inches(0.5)
            else:
                p.paragraph_format.left_indent = Inches(0.75)

            if marker in ['-', '*']:
                bullet_symbol = "• " if level == 0 else "◦ "
                r_symbol = p.add_run(bullet_symbol)
                r_symbol.font.name = "Calibri"
                r_symbol.font.size = Pt(10.5)
                r_symbol.font.bold = True
                r_symbol.font.color.rgb = RGBColor(27, 54, 93)
            else:
                r_num = p.add_run(f"{marker} ")
                r_num.font.name = "Calibri"
                r_num.font.size = Pt(10.5)
                r_num.font.bold = True
                r_num.font.color.rgb = RGBColor(27, 54, 93)

            add_formatted_runs(p, content, base_font_size=10.5, font_name="Calibri")
            i += 1
            continue

        # Regular Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        add_formatted_runs(p, line.strip(), base_font_size=10.5, font_name="Calibri")
        i += 1

    doc.save(docx_filepath)
    print(f"Sucesso! {docx_filepath} foi gerado a partir de {md_filepath}.")

if __name__ == "__main__":
    md_file = os.path.join(os.path.dirname(__file__), "ENTREGA.md")
    docx_file = os.path.join(os.path.dirname(__file__), "ENTREGA.docx")
    parse_markdown_to_docx(md_file, docx_file)
